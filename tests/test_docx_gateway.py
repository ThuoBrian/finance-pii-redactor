"""Tests for the python-docx gateway.

Uses real ``python-docx`` objects (no mocks) to build fixture documents, the
same style as ``test_pdf_gateway.py`` for PyMuPDF.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION

from finance_redactor.domain.entities import Span
from finance_redactor.infrastructure.documents.docx_gateway import PythonDocxDocument


def _build_docx(builder) -> bytes:
    """Build a .docx via ``builder(document)`` and return its bytes."""
    document = Document()
    builder(document)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_blocks_cover_body_table_and_header_paragraphs():
    def build(document):
        document.add_paragraph("A memo about John Doe.")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Vendor: Acme Supplies"
        document.sections[0].header.paragraphs[0].text = "Confidential memo"

    gateway = PythonDocxDocument.open(_build_docx(build))
    texts = [gateway.block_text(i) for i in range(gateway.block_count)]

    assert "A memo about John Doe." in texts
    assert "Vendor: Acme Supplies" in texts
    assert "Confidential memo" in texts


def test_linked_headers_are_not_scanned_twice():
    def build(document):
        document.add_paragraph("body one")
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph("body two")
        document.sections[0].header.paragraphs[0].text = "Shared header"

    gateway = PythonDocxDocument.open(_build_docx(build))
    texts = [gateway.block_text(i) for i in range(gateway.block_count)]

    assert texts.count("Shared header") == 1


def test_replace_block_text_confined_to_one_run():
    def build(document):
        document.add_paragraph("Hello World and more text")

    gateway = PythonDocxDocument.open(_build_docx(build))
    text = gateway.block_text(0)
    start = text.index("World")

    gateway.replace_block_text(0, [(Span(start, start + len("World")), "STF-1")])

    assert gateway.block_text(0) == "Hello STF-1 and more text"


def test_replace_block_text_preserves_formatting_outside_the_span():
    def build(document):
        paragraph = document.add_paragraph("Hello ")
        bold_run = paragraph.add_run("World")
        bold_run.bold = True
        paragraph.add_run(" and more text")

    gateway = PythonDocxDocument.open(_build_docx(build))
    text = gateway.block_text(0)
    start = text.index("World")

    # Span crosses from the plain prefix run into the bold "World" run.
    gateway.replace_block_text(0, [(Span(start - 2, start + len("World")), "PSN-2")])

    runs = gateway._paragraphs[0].runs
    assert runs[-1].text == " and more text"
    assert runs[-1].bold is None
    assert "PSN-2" in "".join(r.text for r in runs)


def test_replace_block_text_handles_multiple_spans_in_one_call():
    def build(document):
        document.add_paragraph("John paid Mary today")

    gateway = PythonDocxDocument.open(_build_docx(build))
    text = gateway.block_text(0)
    john = text.index("John")
    mary = text.index("Mary")

    gateway.replace_block_text(
        0,
        [
            (Span(john, john + 4), "STF-1"),
            (Span(mary, mary + 4), "STF-2"),
        ],
    )

    assert gateway.block_text(0) == "STF-1 paid STF-2 today"


def test_to_bytes_round_trips_replacements():
    def build(document):
        document.add_paragraph("Invoice for John Doe")

    gateway = PythonDocxDocument.open(_build_docx(build))
    text = gateway.block_text(0)
    start = text.index("John Doe")
    gateway.replace_block_text(0, [(Span(start, start + len("John Doe")), "STF-99")])

    reopened = PythonDocxDocument.open(gateway.to_bytes())
    assert "STF-99" in reopened.block_text(0)
    assert "John" not in reopened.block_text(0)
