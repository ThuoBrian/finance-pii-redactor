"""Word (.docx) read/write adapter (python-docx).

Implements the :class:`WordDocument` port. Wraps a ``docx.Document`` and
exposes only the operations the use case needs: enumerating paragraph
"blocks" (body, table cells, headers/footers), reading a block's flattened
text, and splicing pseudonyms back into the underlying runs so unaffected
text keeps its original formatting.
"""

from __future__ import annotations

from io import BytesIO
from typing import cast

from docx import Document as open_docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from finance_redactor.domain.entities import Span


def _table_paragraphs(table: Table) -> list[Paragraph]:
    """Return every paragraph in a table's cells, recursing into nested tables."""
    paragraphs: list[Paragraph] = []
    for row in table.rows:
        for cell in row.cells:
            paragraphs.extend(cell.paragraphs)
            for nested_table in cell.tables:
                paragraphs.extend(_table_paragraphs(nested_table))
    return paragraphs


def _collect_paragraphs(document: Document) -> list[Paragraph]:
    """Enumerate every paragraph in the document: body, tables, headers/footers.

    Headers/footers linked across sections (``is_linked_to_previous``) share
    the same underlying part, so each distinct part is only visited once
    (tracked by ``id(part)``) to avoid scanning/replacing the same text twice.

    Known limitation: text inside text boxes, SmartArt, and embedded objects
    is not part of ``paragraph.runs`` and is not scanned - mirrors the PDF
    flow's "only the selectable text layer is processed" limitation.
    """
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        paragraphs.extend(_table_paragraphs(table))

    seen_part_ids: set[int] = set()
    for section in document.sections:
        for header_or_footer in (section.header, section.footer):
            part_id = id(header_or_footer.part)
            if part_id in seen_part_ids:
                continue
            seen_part_ids.add(part_id)
            paragraphs.extend(header_or_footer.paragraphs)
            for table in header_or_footer.tables:
                paragraphs.extend(_table_paragraphs(table))
    return paragraphs


class PythonDocxDocument:
    """A single open Word document being pseudonymized paragraph by paragraph."""

    def __init__(self, document: Document) -> None:
        """Wrap an already-open python-docx ``Document`` and index its blocks."""
        self._document = document
        self._paragraphs = _collect_paragraphs(document)

    @classmethod
    def open(cls, source: object) -> PythonDocxDocument:
        """Open a .docx from bytes or a readable file-like object."""
        data = source.read() if hasattr(source, "read") else source
        # `source` is deliberately typed as `object` at the port boundary (it
        # may be raw bytes or any file-like upload, e.g. Streamlit's
        # UploadedFile) - the actual runtime contract (bytes in, either way)
        # isn't expressible there without narrowing the port itself.
        return cls(open_docx(BytesIO(cast(bytes, data))))

    @property
    def block_count(self) -> int:
        """Total number of paragraph blocks."""
        return len(self._paragraphs)

    def block_text(self, block_index: int) -> str:
        """Return the flattened text (all runs concatenated) of one block."""
        return "".join(run.text for run in self._paragraphs[block_index].runs)

    def replace_block_text(
        self, block_index: int, replacements: list[tuple[Span, str]]
    ) -> None:
        """Splice each ``(span, pseudonym)`` into the block's runs in place.

        Spans are offsets into ``block_text(block_index)`` (the original,
        pre-replacement text). Replacements are applied right-to-left (like
        the domain's ``apply_replacements``) so an earlier span's offsets stay
        valid while a later one is rewritten. For a span crossing more than
        one run, the pseudonym is inserted once, in the run where the span
        starts (taking that run's formatting); any other run's overlapping
        portion is blanked, and text outside every span keeps its own run and
        formatting untouched.
        """
        runs = self._paragraphs[block_index].runs
        if not runs:
            return
        texts = [run.text for run in runs]
        starts: list[int] = []
        position = 0
        for text in texts:
            starts.append(position)
            position += len(text)
        lengths = [len(text) for text in texts]

        ordered = sorted(replacements, key=lambda item: item[0].start, reverse=True)
        for span, pseudonym in ordered:
            inserted = False
            for i in range(len(runs)):
                run_start = starts[i]
                run_end = run_start + lengths[i]
                overlap_start = max(span.start, run_start)
                overlap_end = min(span.end, run_end)
                if overlap_start >= overlap_end:
                    continue
                local_start = overlap_start - run_start
                local_end = overlap_end - run_start
                replacement_text = pseudonym if not inserted else ""
                texts[i] = (
                    texts[i][:local_start] + replacement_text + texts[i][local_end:]
                )
                inserted = True

        for run, text in zip(runs, texts):
            if run.text != text:
                run.text = text

    def to_bytes(self) -> bytes:
        """Render the pseudonymized document to bytes."""
        output = BytesIO()
        self._document.save(output)
        return output.getvalue()

    def close(self) -> None:
        """Release underlying resources (no-op: python-docx holds nothing open)."""
